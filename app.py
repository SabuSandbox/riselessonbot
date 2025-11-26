# app.py
import os
import re
import json
import tempfile
import nltk
import requests
from flask import Flask, request, jsonify
from docx import Document
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
from duckduckgo_search import ddg
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.text_rank import TextRankSummarizer

# ---------------- NLTK setup (ensure punkt) ----------------
NLTK_DATA_DIR = os.environ.get("NLTK_DATA_DIR", "/opt/render/nltk_data")
os.makedirs(NLTK_DATA_DIR, exist_ok=True)
if NLTK_DATA_DIR not in nltk.data.path:
    nltk.data.path.append(NLTK_DATA_DIR)
try:
    nltk.data.find("tokenizers/punkt")
except LookupError:
    try:
        nltk.download("punkt", download_dir=NLTK_DATA_DIR, quiet=True)
    except Exception as e:
        print("NLTK download failed:", e)

# ---------------- Config ----------------
TELEGRAM_TOKEN = os.environ.get("TELEGRAM_TOKEN")
if not TELEGRAM_TOKEN:
    raise RuntimeError("TELEGRAM_TOKEN environment variable is required")

DEFAULT_TEMPLATE_PATH = os.environ.get("DEFAULT_TEMPLATE_PATH", "./Sample Lesson Plan.docx")
PORT = int(os.environ.get("PORT", 5000))
ADMIN_ID = int(os.environ.get("ADMIN_ID", "7925575742"))
TARGET_USER_ID_ENV = os.environ.get("TARGET_USER_ID")  # optional

BASE_TELEGRAM_URL = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}"

app = Flask(__name__)

# In-memory session storage (ephemeral)
SESS = {}           # chat_id -> {state, tmp, template_path}
RUNTIME_TARGET = {} # runtime override for admin target

# ---------------- Telegram helpers ----------------
def telegram_api(method, params=None, files=None, json_payload=None):
    url = f"{BASE_TELEGRAM_URL}/{method}"
    try:
        if files:
            return requests.post(url, params=params, files=files, timeout=30)
        if json_payload:
            return requests.post(url, json=json_payload, timeout=30)
        return requests.post(url, data=params, timeout=30)
    except Exception as e:
        print("telegram_api error:", e)
        raise

def send_message(chat_id, text, reply_markup=None):
    payload = {"chat_id": chat_id, "text": text}
    if reply_markup:
        payload["reply_markup"] = json.dumps(reply_markup)
    try:
        telegram_api("sendMessage", params=payload)
    except Exception as e:
        print("send_message error:", e)

def download_file(file_id, dest_path):
    r = telegram_api("getFile", params={"file_id": file_id})
    r.raise_for_status()
    data = r.json()
    file_path = data["result"]["file_path"]
    file_url = f"https://api.telegram.org/file/bot{TELEGRAM_TOKEN}/{file_path}"
    r2 = requests.get(file_url, timeout=60)
    r2.raise_for_status()
    with open(dest_path, "wb") as f:
        f.write(r2.content)
    return dest_path

# ---------------- Extraction helpers ----------------
def extract_text_from_pdf(path):
    """Extract plain text from PDF using PyPDF2 (best-effort)."""
    text_parts = []
    try:
        reader = PdfReader(path)
        for page in reader.pages:
            try:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
            except Exception:
                continue
    except Exception as e:
        print("PDF read error:", e)
        raise
    return "\n".join(text_parts)

def extract_text_from_url(url, max_chars=20000):
    """Lightweight extractor: prefer <article>, otherwise join large <p> blocks."""
    try:
        r = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
        r.raise_for_status()
    except Exception as e:
        print("fetch error for", url, ":", e)
        return ""
    soup = BeautifulSoup(r.text, "html.parser")
    article = soup.find("article")
    if article:
        text = article.get_text(separator="\n")
    else:
        ps = soup.find_all("p")
        filtered = []
        for p in ps:
            t = p.get_text(strip=True)
            if not t:
                continue
            if len(t) < 30:  # skip very short fragments
                continue
            filtered.append(t)
        text = "\n\n".join(filtered)
    if not text or len(text.strip()) < 100:
        title = soup.title.string.strip() if soup.title and soup.title.string else ""
        desc_tag = soup.find("meta", attrs={"name":"description"}) or soup.find("meta", attrs={"property":"og:description"})
        meta = desc_tag.get("content").strip() if desc_tag and desc_tag.get("content") else ""
        text = (title + "\n" + meta).strip()
    return (text or "")[:max_chars]

# ---------------- Summarization and heuristics ----------------
def summarize_text(text, sentences_count=6):
    if not text:
        return ""
    try:
        parser = PlaintextParser.from_string(text, Tokenizer("english"))
        summarizer = TextRankSummarizer()
        summary_sentences = summarizer(parser.document, sentences_count)
        return "\n".join(str(s) for s in summary_sentences)
    except Exception as e:
        print("summarize_text error:", e)
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        return "\n".join(lines[:sentences_count]) if lines else (text[:1000] if text else "")

def extract_objectives_from_text(text, max_points=5):
    candidates = []
    for sent in (text or "").split("."):
        s = sent.strip()
        if not s:
            continue
        sl = s.lower()
        if any(k in sl for k in ("able to", "will", "understand", "learn", "identify", "describe", "students will")):
            candidates.append(s.strip())
    if candidates:
        return "\n".join(f"• {c}" for c in candidates[:max_points])
    summ = summarize_text(text, sentences_count=max_points)
    if summ:
        return "\n".join(f"• {s.strip()}" for s in summ.split("\n") if s.strip())
    return "• Students will be able to ..."

def generate_activities(text, max_items=4):
    return (
        "1. Read the summary and discuss key terms.\n"
        "2. Small-group activity: identify examples from the text.\n"
        "3. Hands-on/demo (if applicable): follow the experiment steps.\n"
        "4. Exit ticket: one short question to assess learning."
    )

def generate_assessment_questions(text, max_q=4):
    summary = summarize_text(text, sentences_count=4)
    qs = []
    for i, s in enumerate(summary.split("\n")[:max_q], start=1):
        s = s.strip().rstrip(".")
        if len(s.split()) < 3:
            continue
        qs.append(f"Q{i}. Explain: {s}?")
    if not qs:
        qs = ["Q1. What is the main idea of the chapter?", "Q2. List two key points."]
    return "\n".join(qs)

# ---------------- Section extractor heuristics ----------------
def extract_section(text, names, window_chars=800):
    """
    Naive extractor: search for heading and return the following text.
    names: a regex alternation like "Resource|Resources|Materials"
    """
    if not text:
        return ""
    pattern = rf'({names})\s*[:\-\n]\s*(.*?)(\n[A-Z][^\n]{{0,80}}:|\Z)'
    m = re.search(pattern, text, flags=re.I | re.S)
    if m:
        return m.group(2).strip()[:window_chars]
    return ""

# ---------------- DOCX replacement helpers for human-readable labels ----------------
def _set_paragraph_text(paragraph, new_text):
    for i in range(len(paragraph.runs) - 1, -1, -1):
        paragraph.runs[i].clear()
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)

def _replace_in_paragraph_by_label(paragraph, label, replacement):
    ptext = paragraph.text or ""
    lower = ptext.lower()
    lab = label.lower()
    if lab not in lower:
        return False
    # bracketed placeholder
    bracket_match = re.search(r'\[([^\]]*)\]', ptext)
    if bracket_match:
        new_ptext = re.sub(r'\[([^\]]*)\]', replacement, ptext, count=1)
        _set_paragraph_text(paragraph, new_ptext)
        return True
    # colon present -> replace after colon
    if ':' in ptext:
        parts = ptext.split(':', 1)
        left = parts[0].rstrip()
        new_ptext = f"{left}: {replacement}"
        _set_paragraph_text(paragraph, new_ptext)
        return True
    # fallback -> overwrite paragraph
    _set_paragraph_text(paragraph, replacement)
    return True

def _replace_in_table(table, label, replacement):
    done = False
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if _replace_in_paragraph_by_label(p, label, replacement):
                    done = True
    return done

def _replace_in_headers_footers(doc, label, replacement):
    done = False
    try:
        for section in doc.sections:
            hdr = section.header
            for p in hdr.paragraphs:
                if _replace_in_paragraph_by_label(p, label, replacement):
                    done = True
            for t in hdr.tables:
                if _replace_in_table(t, label, replacement):
                    done = True
            ftr = section.footer
            for p in ftr.paragraphs:
                if _replace_in_paragraph_by_label(p, label, replacement):
                    done = True
            for t in ftr.tables:
                if _replace_in_table(t, label, replacement):
                    done = True
    except Exception:
        pass
    return done

def _replace_in_doc(doc, label, replacement):
    for p in doc.paragraphs:
        if _replace_in_paragraph_by_label(p, label, replacement):
            return True
    for t in doc.tables:
        if _replace_in_table(t, label, replacement):
            return True
    if _replace_in_headers_footers(doc, label, replacement):
        return True
    return False

# ---------------- Main fill function for your template ----------------
def fill_template_and_send_bracketed(chat_id, mapping):
    """
    mapping keys (example):
      lesson_title, grade, subject, teacher_name, date,
      objectives, resources, outline, assessment, homework, conclusion, note
    """
    template_path = SESS.get(chat_id, {}).get("template_path") or DEFAULT_TEMPLATE_PATH
    if not os.path.exists(template_path):
        send_message(chat_id, "Template not found on server. Please upload a .docx template or set DEFAULT_TEMPLATE_PATH.")
        return

    # load document
    doc = Document(template_path)

    # label variants dictionary (tune to your template's wording if needed)
    label_variants = {
        "lesson_title": ["lesson plant", "lesson plan", "lesson title", "chapter name"],
        "grade": ["grade"],
        "subject": ["subject"],
        "teacher_name": ["teacher name", "teacher"],
        "date": ["date"],
        "objectives": ["lesson objectives", "lesson objective", "objectives"],
        "resources": ["resource needed", "resources", "materials"],
        "outline": ["lesson outline", "lesson outline:"],
        "assessment": ["assessment and evaluation", "assessment", "evaluation"],
        "homework": ["homework/extension activity", "homework", "extension activity", "assignment"],
        "conclusion": ["conclusion", "summary"],
        "note": ["note for teacher", "note", "notes"]
    }

    # run replacements
    for canonical_key, variants in label_variants.items():
        if canonical_key not in mapping:
            continue
        replacement_text = mapping.get(canonical_key) or ""
        for var in variants:
            if _replace_in_doc(doc, var, replacement_text):
                break

    # fallback: replace any remaining bracket tokens with leftover mapping values
    leftover_values = [v for k, v in mapping.items() if v]
    if leftover_values:
        bracket_pattern = re.compile(r'\[([^\]]+)\]')
        for p in doc.paragraphs:
            def repl_fn(m):
                return leftover_values.pop(0) if leftover_values else m.group(0)
            new_text = bracket_pattern.sub(repl_fn, p.text)
            if new_text != p.text:
                _set_paragraph_text(p, new_text)

    # save and send
    out_tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    out_path = out_tmp.name
    doc.save(out_path)
    try:
        with open(out_path, "rb") as f:
            files = {"document": (os.path.basename(out_path), f)}
            telegram_api("sendDocument", params={"chat_id": chat_id}, files=files)
        send_message(chat_id, "Lesson plan generated ✅")
    except Exception as e:
        print("sendDocument error:", e)
        send_message(chat_id, f"Failed to send generated file: {e}")

# ---------------- Admin helpers ----------------
def is_admin(chat_id):
    try:
        return int(chat_id) == int(ADMIN_ID)
    except Exception:
        return False

def get_current_target():
    rt = RUNTIME_TARGET.get("target")
    if rt:
        return str(rt)
    if TARGET_USER_ID_ENV:
        return str(TARGET_USER_ID_ENV)
    return None

# ---------------- Webhook & conversation flows ----------------
@app.route("/health", methods=["GET"])
def health():
    return jsonify({"ok": True})

@app.route("/webhook", methods=["POST"])
def webhook():
    update = request.get_json(force=True)
    if not update:
        return jsonify({"ok": True})
    # only process messages (safe)
    if "message" not in update:
        return jsonify({"ok": True})
    msg = update["message"]
    chat_id = msg["chat"]["id"]
    SESS.setdefault(chat_id, {"state": "idle", "tmp": {}, "template_path": None})

    # ---------- Quick admin CLI commands (text-based) ----------
    if "text" in msg:
        text = msg["text"].strip()
        # /admin opens menu (admin only)
        if text.startswith("/admin"):
            if not is_admin(chat_id):
                send_message(chat_id, "Unauthorized. Only admin can use this command.")
                return jsonify({"ok": True})
            kb = {"keyboard":[["Send Message to Target"],["Show Target"],["Set Target"],["Set Template Path"],["Exit Admin"]],"one_time_keyboard":True,"resize_keyboard":True}
            send_message(chat_id, "Admin menu — choose an action:", reply_markup=kb)
            SESS[chat_id]["state"] = "admin_menu"
            return jsonify({"ok": True})

        # /settarget <id>
        if text.startswith("/settarget"):
            if not is_admin(chat_id):
                send_message(chat_id, "Unauthorized.")
                return jsonify({"ok": True})
            parts = text.split(maxsplit=1)
            if len(parts) == 2 and parts[1].strip().isdigit():
                RUNTIME_TARGET["target"] = parts[1].strip()
                send_message(chat_id, f"Runtime target set to: {RUNTIME_TARGET['target']}")
            else:
                send_message(chat_id, "Usage: /settarget <chat_id>")
            return jsonify({"ok": True})

        # /showtarget
        if text.startswith("/showtarget"):
            if not is_admin(chat_id):
                send_message(chat_id, "Unauthorized.")
                return jsonify({"ok": True})
            cur = get_current_target()
            send_message(chat_id, f"Current target: {cur or 'None'}")
            return jsonify({"ok": True})

        # /sendtarget <message>
        if text.startswith("/sendtarget"):
            if not is_admin(chat_id):
                send_message(chat_id, "Unauthorized.")
                return jsonify({"ok": True})
            parts = text.split(maxsplit=1)
            if len(parts) != 2:
                send_message(chat_id, "Usage: /sendtarget <message>")
                return jsonify({"ok": True})
            target = get_current_target()
            if not target:
                send_message(chat_id, "No target set. Use /settarget or set TARGET_USER_ID env var.")
                return jsonify({"ok": True})
            send_message(target, parts[1].strip())
            send_message(chat_id, f"Message sent to {target}")
            return jsonify({"ok": True})

    # ---------- Admin menu handling ----------
    if "text" in msg and SESS[chat_id].get("state") == "admin_menu" and is_admin(chat_id):
        choice = msg["text"].strip()
        if choice == "Send Message to Target":
            send_message(chat_id, "Please send the message you want to forward to the target (single message).")
            SESS[chat_id]["state"] = "admin_send_message"
            return jsonify({"ok": True})
        if choice == "Show Target":
            cur = get_current_target()
            send_message(chat_id, f"Current target: {cur or 'None'}")
            SESS[chat_id]["state"] = "admin_menu"
            return jsonify({"ok": True})
        if choice == "Set Target":
            send_message(chat_id, "Send the chat_id to set as runtime target (digits only).")
            SESS[chat_id]["state"] = "admin_set_target"
            return jsonify({"ok": True})
        if choice == "Set Template Path":
            send_message(chat_id, "Send the absolute path (inside container) to set as default template for admin (e.g. ./Sample Lesson Plan.docx).")
            SESS[chat_id]["state"] = "admin_set_template"
            return jsonify({"ok": True})
        if choice == "Exit Admin":
            send_message(chat_id, "Exiting admin menu.")
            SESS[chat_id]["state"] = "idle"
            return jsonify({"ok": True})
        send_message(chat_id, "Unknown admin choice.")
        SESS[chat_id]["state"] = "idle"
        return jsonify({"ok": True})

    # ---------- Admin follow-ups ----------
    if "text" in msg and is_admin(chat_id):
        state = SESS[chat_id].get("state")
        if state == "admin_send_message":
            message_to_send = msg["text"].strip()
            target = get_current_target()
            if not target:
                send_message(chat_id, "No target set. Use Set Target option or /settarget <chat_id>.")
                SESS[chat_id]["state"] = "admin_menu"
                return jsonify({"ok": True})
            send_message(target, message_to_send)
            send_message(chat_id, f"Message sent to {target}")
            SESS[chat_id]["state"] = "admin_menu"
            return jsonify({"ok": True})
        if state == "admin_set_target":
            cand = msg["text"].strip()
            if cand.isdigit():
                RUNTIME_TARGET["target"] = cand
                send_message(chat_id, f"Runtime target set to {cand}")
            else:
                send_message(chat_id, "Invalid chat_id. Digits only.")
            SESS[chat_id]["state"] = "admin_menu"
            return jsonify({"ok": True})
        if state == "admin_set_template":
            cand = msg["text"].strip()
            if os.path.exists(cand):
                SESS[chat_id]["template_path"] = cand
                send_message(chat_id, f"Admin template path set to: {cand}")
            else:
                send_message(chat_id, f"Path does not exist in container: {cand}")
            SESS[chat_id]["state"] = "admin_menu"
            return jsonify({"ok": True})

    # ---------- Normal user flows ----------
    SESS.setdefault(chat_id, {"state": "idle", "tmp": {}, "template_path": None})

    # entry command /hi_rise
    if "text" in msg and msg["text"].strip().lower() == "/hi_rise":
        kb = {"keyboard":[["Upload PDF"],["Paste Text"],["Ask Bot to Find Lesson"]],"one_time_keyboard":True,"resize_keyboard":True}
        send_message(chat_id, "Hi! For which lesson shall we create a lesson plan today? Choose how you'd like to provide the lesson:", reply_markup=kb)
        SESS[chat_id]["state"] = "idle"
        SESS[chat_id]["tmp"] = {}
        return jsonify({"ok": True})

    # user options when idle
    if "text" in msg and SESS[chat_id]["state"] == "idle":
        txt = msg["text"].strip()
        if txt == "Upload PDF":
            SESS[chat_id]["state"] = "await_pdf"
            send_message(chat_id, "Please upload the lesson PDF as a document now (attach as Telegram Document).")
            return jsonify({"ok": True})
        if txt == "Paste Text":
            SESS[chat_id]["state"] = "await_text"
            send_message(chat_id, "Please paste the chapter text now.")
            return jsonify({"ok": True})
        if txt == "Ask Bot to Find Lesson":
            SESS[chat_id]["state"] = "await_grade"
            send_message(chat_id, "Okay — which Grade? (e.g., Grade 6)")
            return jsonify({"ok": True})
        if len(txt) > 120:
            send_message(chat_id, "I detected pasted text — reply 'Yes' to confirm generation.")
            SESS[chat_id]["state"] = "confirm_from_text"
            SESS[chat_id]["tmp"]["text_candidate"] = txt
            return jsonify({"ok": True})

    # handle document uploads (PDFs or .docx template)
    if "document" in msg:
        doc = msg["document"]
        fname = doc.get("file_name", "file")
        file_id = doc["file_id"]
        tmpdir = tempfile.mkdtemp()
        local_path = os.path.join(tmpdir, fname)
        try:
            download_file(file_id, local_path)
        except Exception as e:
            send_message(chat_id, f"Failed to download file: {e}")
            SESS[chat_id]["state"] = "idle"
            return jsonify({"ok": True})

        # .docx -> template
        if fname.lower().endswith(".docx"):
            SESS[chat_id]["template_path"] = local_path
            send_message(chat_id, "Template uploaded and saved for your session. Now upload PDF, paste text, or use /hi_rise to start again.")
            SESS[chat_id]["state"] = "idle"
            return jsonify({"ok": True})

        # pdf handling when awaiting pdf
        if fname.lower().endswith(".pdf") and SESS[chat_id]["state"] == "await_pdf":
            try:
                pdf_text = extract_text_from_pdf(local_path)
            except Exception as e:
                send_message(chat_id, f"PDF extraction failed: {e}")
                SESS[chat_id]["state"] = "idle"
                return jsonify({"ok": True})

            send_message(chat_id, "PDF received. Generating lesson plan...")
            # extract fields
            title = SESS[chat_id]["tmp"].get("chapter_title") or "Lesson"
            summary = summarize_text(pdf_text, sentences_count=6)
            objectives = extract_objectives_from_text(pdf_text)
            activities = generate_activities(pdf_text)
            assessment = generate_assessment_questions(pdf_text)
            resources = extract_section(pdf_text, "Resource|Resources|Materials")
            outline = summarize_text(pdf_text, sentences_count=8)
            homework = extract_section(pdf_text, "Homework|Extension Activity|Assignment")
            conclusion = extract_section(pdf_text, "Conclusion|Summary|Summing up")
            note = extract_section(pdf_text, "Note for Teacher|Teacher Note|Notes")

            mapping = {
                "lesson_title": title,
                "grade": SESS[chat_id]["tmp"].get("grade",""),
                "subject": SESS[chat_id]["tmp"].get("subject",""),
                "teacher_name": SESS[chat_id]["tmp"].get("teacher",""),
                "date": SESS[chat_id]["tmp"].get("date",""),
                "objectives": objectives,
                "resources": resources,
                "outline": outline,
                "assessment": assessment,
                "homework": homework,
                "conclusion": conclusion,
                "note": note
            }
            fill_template_and_send_bracketed(chat_id, mapping)
            SESS[chat_id]["state"] = "idle"
            return jsonify({"ok": True})

        send_message(chat_id, "Document received. If this is a PDF for lesson content, please choose 'Upload PDF' first. If this is a .docx template, it has been saved.")
        SESS[chat_id]["state"] = "idle"
        return jsonify({"ok": True})

    # photos (OCR not enabled)
    if "photo" in msg:
        send_message(chat_id, "Photo received. OCR is not enabled in this deployment. Please upload a PDF or paste text.")
        return jsonify({"ok": True})

    # plain text during flows
    if "text" in msg:
        txt = msg["text"].strip()
        state = SESS[chat_id]["state"]

        if state == "confirm_from_text":
            if txt.lower() in ("yes","y"):
                candidate = SESS[chat_id]["tmp"].get("text_candidate","")
                SESS[chat_id]["state"] = "idle"
                send_message(chat_id, "Generating lesson plan from pasted text...")
                summary = summarize_text(candidate, sentences_count=6)
                objectives = extract_objectives_from_text(candidate)
                activities = generate_activities(candidate)
                assessment = generate_assessment_questions(candidate)
                resources = extract_section(candidate, "Resource|Resources|Materials")
                outline = summarize_text(candidate, sentences_count=8)
                homework = extract_section(candidate, "Homework|Extension Activity|Assignment")
                conclusion = extract_section(candidate, "Conclusion|Summary|Summing up")
                note = extract_section(candidate, "Note for Teacher|Teacher Note|Notes")
                mapping = {
                    "lesson_title": SESS[chat_id]["tmp"].get("chapter_title","Pasted Lesson"),
                    "grade": SESS[chat_id]["tmp"].get("grade",""),
                    "subject": SESS[chat_id]["tmp"].get("subject",""),
                    "teacher_name": SESS[chat_id]["tmp"].get("teacher",""),
                    "date": SESS[chat_id]["tmp"].get("date",""),
                    "objectives": objectives,
                    "resources": resources,
                    "outline": outline,
                    "assessment": assessment,
                    "homework": homework,
                    "conclusion": conclusion,
                    "note": note
                }
                fill_template_and_send_bracketed(chat_id, mapping)
                return jsonify({"ok": True})
            else:
                SESS[chat_id]["state"] = "idle"
                send_message(chat_id, "Cancelled. Send /hi_rise to begin again.")
                return jsonify({"ok": True})

        if state == "await_text":
            text = txt
            SESS[chat_id]["state"] = "idle"
            send_message(chat_id, "Generating lesson plan from pasted text...")
            summary = summarize_text(text, sentences_count=6)
            objectives = extract_objectives_from_text(text)
            activities = generate_activities(text)
            assessment = generate_assessment_questions(text)
            resources = extract_section(text, "Resource|Resources|Materials")
            outline = summarize_text(text, sentences_count=8)
            homework = extract_section(text, "Homework|Extension Activity|Assignment")
            conclusion = extract_section(text, "Conclusion|Summary|Summing up")
            note = extract_section(text, "Note for Teacher|Teacher Note|Notes")
            mapping = {
                "lesson_title": SESS[chat_id]["tmp"].get("chapter_title","Pasted Lesson"),
                "grade": SESS[chat_id]["tmp"].get("grade",""),
                "subject": SESS[chat_id]["tmp"].get("subject",""),
                "teacher_name": SESS[chat_id]["tmp"].get("teacher",""),
                "date": SESS[chat_id]["tmp"].get("date",""),
                "objectives": objectives,
                "resources": resources,
                "outline": outline,
                "assessment": assessment,
                "homework": homework,
                "conclusion": conclusion,
                "note": note
            }
            fill_template_and_send_bracketed(chat_id, mapping)
            return jsonify({"ok": True})

        # Ask-bot-to-find flow: grade -> subject -> chapter
        if state == "await_grade":
            SESS[chat_id]["tmp"] = {"grade": txt}
            SESS[chat_id]["state"] = "await_subject"
            send_message(chat_id, "Which Subject? (e.g., Mathematics, Science, English)")
            return jsonify({"ok": True})

        if state == "await_subject":
            SESS[chat_id]["tmp"]["subject"] = txt
            SESS[chat_id]["state"] = "await_chapter"
            send_message(chat_id, "Which Chapter name or number should I search for?")
            return jsonify({"ok": True})

        if state == "await_chapter":
            SESS[chat_id]["tmp"]["chapter"] = txt
            grade = SESS[chat_id]["tmp"].get("grade")
            subject = SESS[chat_id]["tmp"].get("subject")
            chapter = SESS[chat_id]["tmp"].get("chapter")
            query = f"{grade} {subject} {chapter} summary lesson"
            send_message(chat_id, f"Searching web for: {query}")
            SESS[chat_id]["state"] = "idle"
            try:
                hits = ddg(query, max_results=5) or []
            except Exception as e:
                print("ddg error:", e)
                hits = []
            combined_texts = []
            references = []
            for h in hits[:3]:
                url = h.get("href") or h.get("url")
                title = h.get("title") or url
                snippet = h.get("body") or h.get("snippet") or ""
                txt_extracted = extract_text_from_url(url) if url else snippet
                if txt_extracted:
                    combined_texts.append(txt_extracted)
                references.append(f"{title} — {url}")
            big_text = "\n\n".join(combined_texts) or " ".join([h.get("body","") or h.get("snippet","") for h in hits]) or chapter
            summary = summarize_text(big_text, sentences_count=6)
            objectives = extract_objectives_from_text(big_text)
            activities = generate_activities(big_text)
            assessment = generate_assessment_questions(big_text)
            resources = extract_section(big_text, "Resource|Resources|Materials")
            outline = summarize_text(big_text, sentences_count=8)
            homework = extract_section(big_text, "Homework|Extension Activity|Assignment")
            conclusion = extract_section(big_text, "Conclusion|Summary|Summing up")
            note = extract_section(big_text, "Note for Teacher|Teacher Note|Notes")
            mapping = {
                "lesson_title": chapter,
                "grade": grade or "",
                "subject": subject or "",
                "teacher_name": "",
                "date": "",
                "objectives": objectives,
                "resources": resources,
                "outline": outline,
                "assessment": assessment,
                "homework": homework,
                "conclusion": conclusion,
                "note": note
            }
            fill_template_and_send_bracketed(chat_id, mapping)
            return jsonify({"ok": True})

        # fallback help
        send_message(chat_id, "Send /hi_rise to start the lesson-plan flow, or upload a .docx template.")
        return jsonify({"ok": True})

    return jsonify({"ok": True})

# ---------------- run app locally ----------------
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=PORT)
